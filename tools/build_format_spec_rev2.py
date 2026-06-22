from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path("artifacts") / "Format specification - rev2.docx"

COLOR_H1 = RGBColor(0x2E, 0x74, 0xB5)
COLOR_H3 = RGBColor(0x1F, 0x4D, 0x78)
COLOR_TEXT = RGBColor(0x22, 0x22, 0x22)
COLOR_MUTED = RGBColor(0x66, 0x66, 0x66)
COLOR_TABLE_HEADER = "E8EEF5"
COLOR_TABLE_ALT = "F7F9FC"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, name="Calibri", size=11, bold=False, italic=False, color=COLOR_TEXT):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def style_base_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = COLOR_TEXT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, COLOR_H1, 18, 10),
        ("Heading 2", 13, COLOR_H1, 14, 7),
        ("Heading 3", 12, COLOR_H3, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15


def add_title_block(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Format specification - rev2")
    set_run_font(run, size=22, bold=True, color=COLOR_H3)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(
        "Technical specification for encrypted EMG and IMU recording. "
        "This rev2 document updates the original design note dated 2026-04-23 "
        "to reflect the current implementation decisions and the accepted "
        "format simplifications."
    )
    set_run_font(run, size=10.5, color=COLOR_MUTED)


def add_label_paragraph(doc: Document, label: str, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    lead = p.add_run(label + ": ")
    set_run_font(lead, bold=True)
    body = p.add_run(text)
    set_run_font(body)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_layout(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()

            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))

            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, value in CELL_MARGIN_DXA.items():
                margin = tc_mar.find(qn(f"w:{side}"))
                if margin is None:
                    margin = OxmlElement(f"w:{side}")
                    tc_mar.append(margin)
                margin.set(qn("w:w"), str(value))
                margin.set(qn("w:type"), "dxa")


def format_table_text(cell, text, *, bold=False, font_size=9.5, color=COLOR_TEXT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    set_run_font(run, size=font_size, bold=bold, color=color)


def add_table(doc: Document, headers, rows, widths_dxa):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    set_table_layout(table, widths_dxa)
    set_repeat_table_header(table.rows[0])

    for idx, header in enumerate(headers):
        format_table_text(table.rows[0].cells[idx], header, bold=True)
        set_cell_shading(table.rows[0].cells[idx], COLOR_TABLE_HEADER)

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            format_table_text(table.rows[r_idx].cells[c_idx], str(value))
            if r_idx % 2 == 0:
                set_cell_shading(table.rows[r_idx].cells[c_idx], COLOR_TABLE_ALT)
    return table


def revision_rows():
    return [
        (
            "File structure",
            "Remove the optional final summary record and define the file as header plus repeated batches only.",
            "Recording can stop at any time because of power loss or deliberate shutdown. Batch trailers are the only reliable completion boundary.",
        ),
        (
            "Payload dispatch",
            "Use payload_type as the sole chunk-layout discriminator and leave header byte 39 reserved.",
            "This removes duplicated dispatch logic. Readers inspect one field to decide how to parse the payload.",
        ),
        (
            "Lead-off handling",
            "Do not store a per-sample status byte. Detect lead-off in software per 1-second chunk and report a per-batch ch1_leadoff_chunks summary (plus the ch1_diff_abs_sum level metric) in the control CSV.",
            "This preserves useful signal-quality observability without increasing the binary payload by 1000 bytes per second.",
        ),
        (
            "Nonce scheme",
            "Fix the nonce scheme to an entropy-seeded 8-byte session salt plus the 4-byte batch index.",
            "The earlier RTC and session-id approach risks nonce reuse under a fixed key. Rev2 makes uniqueness independent of RTC state.",
        ),
        (
            "Cipher definition",
            "State AES-256-GCM as the normative cipher for the current format.",
            "The current STM32 firmware uses the hardware AES path, so the specification should describe the implemented algorithm rather than alternatives.",
        ),
        (
            "Current filenames",
            "Document the current SNNNNNNN.EMX and SNNNNNNN.CSV filenames while noting that longer names remain a future option.",
            "The emitted files must match FatFS 8.3 constraints today, but naming does not change the on-disk byte format.",
        ),
    ]


FILE_HEADER_ROWS = [
    ("0", "4", "magic", '"EMGX"'),
    ("4", "1", "version", "1"),
    ("5", "1", "header_size", "64"),
    ("6", "1", "payload_type", "2 = EMG ch1 + IMU chunk"),
    ("7", "1", "batch_duration_s", "10"),
    ("8", "4", "session_id", "u32, matches the filename"),
    ("12", "12", "device_uid", "STM32 96-bit unique device ID"),
    ("24", "6", "start_time_bcd", "RTC BCD: sec,min,hr,date,mon,yr"),
    ("30", "2", "emg_rate_hz", "1000"),
    ("32", "2", "imu_rate_hz", "100"),
    ("34", "1", "cipher_id", "1 = AES-256-GCM"),
    ("35", "1", "key_id", "From recording.key"),
    ("36", "1", "key_version", "From recording.key"),
    ("37", "1", "nonce_scheme", "1 = entropy-seeded random salt + batch counter"),
    ("38", "1", "emg_pga_gain", "ADS1292R CH1 PGA gain (4)"),
    ("39", "1", "reserved", "0"),
    ("40", "8", "nonce_prefix", "Per-session nonce salt"),
    (
        "48",
        "10",
        "ads_regs",
        "ADS1292R snapshot: ID, CONFIG1, CONFIG2, LOFF, CH1SET, CH2SET, RLDSENS, LOFFSENS, RESP1, RESP2",
    ),
    ("58", "2", "emg_ref_mv", "2420"),
    ("60", "4", "reserved", "0"),
]


BATCH_HEADER_ROWS = [
    ("0", "4", "marker", '"BATC"'),
    ("4", "4", "batch_index", "u32, from 0"),
    ("8", "6", "start_time_bcd", "RTC BCD at batch start"),
    ("14", "1", "batch_duration_s", "10"),
    ("15", "1", "reserved", "0"),
    ("16", "4", "emg_sample_count", "u32"),
    ("20", "4", "imu_sample_count", "u32"),
    ("24", "4", "plaintext_payload_bytes", "u32"),
    ("28", "4", "ciphertext_payload_bytes", "u32, equal to plaintext size"),
    ("32", "12", "nonce", "nonce_prefix (8) plus batch_index as u32 little-endian (4)"),
    ("44", "4", "reserved", "0"),
]


PAYLOAD_ROWS = [
    ("ch1", "int32 LE", "1000", "4000", "ADS1292R channel 1, 24-bit signed value sign-extended to int32"),
    ("imu", "ImuSample", "100", "2400", "100 Hz IMU slots derived from the EMG sample clock"),
]


IMU_ROWS = [
    ("0", "int16 x3", "ax, ay, az"),
    ("6", "int16 x3", "gx, gy, gz"),
    ("12", "uint32 x3", "mx, my, mz"),
]


CSV_ROWS = [
    ("session_id", "Session identifier, matches the filename"),
    ("batch_index", "Zero-based batch counter"),
    ("start_timestamp", "Batch start time in BCD YYMMDDhhmmss"),
    ("end_timestamp", "Batch end time in BCD YYMMDDhhmmss"),
    ("emg_sample_count", "Nominally 10000 for a full 10-second batch"),
    ("imu_sample_count", "Nominally 1000 for a full 10-second batch"),
    ("payload_bytes", "Nominally 64000 for a full batch; AES-GCM ciphertext is the same size"),
    ("write_time_ms", "Time spent in f_write and f_sync for the batch"),
    ("crc32_plaintext", "CRC32 of the plaintext payload"),
    ("dropped_samples", "Samples lost because the acquisition ring overflowed"),
    ("temperature_c", "Optional IMU temperature field"),
    ("error_flags", "Hex bit field: WRITE, SYNC, RTC, DROPPED, AES, PARTIAL"),
    ("storage_status", "OK or ERR"),
    ("ch1_saturated_samples", "Count of CH1 samples railed near +/- full-scale"),
    ("ch1_flatline_chunks / ch1_leadoff_chunks / ch1_baseline_drift_chunks",
     "Mutually-exclusive per-1s-chunk state counts (stuck/flat -> disconnected -> drifting)"),
    ("ch1_diff_abs_sum_min / _med / _max",
     "Per-batch min/median/max of each chunk's sum-of-abs-sample-differences (level metric)"),
    # NOTE: FORMAT.md SS8 is authoritative for the control-CSV columns; keep this in sync.
]


def build_document():
    doc = Document()
    style_base_document(doc)
    doc.core_properties.title = "Format specification - rev2"
    doc.core_properties.subject = "AstroMoWe EMGX recording format"
    doc.core_properties.author = "Codex"

    add_title_block(doc)

    doc.add_paragraph("Revision summary", style="Heading 1")
    add_label_paragraph(
        doc,
        "Status",
        "This rev2 document is intended to replace the original prose design note "
        "for current implementation work. The reasoning below is normative where it "
        "changes parser or writer behavior.",
    )
    add_table(
        doc,
        ["Area", "Rev2 decision", "Reasoning"],
        revision_rows(),
        [1800, 3780, 3780],
    )

    doc.add_paragraph("1. Purpose and scope", style="Heading 1")
    doc.add_paragraph(
        "The recording format is intended for efficient and secure storage of EMG "
        "and IMU data acquired by the AstroMoWe embedded device. The format is "
        "optimized for real-time logging on STM32L462-class hardware while remaining "
        "simple to parse on desktop tools."
    )
    doc.add_paragraph(
        "Rev2 defines the current implemented format and removes design-note "
        "ambiguities that are no longer useful for interoperable readers and writers."
    )

    doc.add_paragraph("2. Output files", style="Heading 1")
    doc.add_paragraph(
        "Each recording session produces two coordinated files that share the same "
        "session identifier."
    )
    add_table(
        doc,
        ["File", "Current implementation", "Purpose"],
        [
            ("Primary binary file", "SNNNNNNN.EMX", "Encrypted measurement data"),
            ("Secondary control file", "SNNNNNNN.CSV", "Plaintext operational metadata"),
        ],
        [1800, 2160, 5400],
    )
    add_label_paragraph(
        doc,
        "Naming note",
        "The current firmware uses 8.3 filenames because FatFS long filename support "
        "is disabled. Longer SESSION_NNNNNNN names remain an optional future change "
        "and do not affect the byte format.",
    )

    doc.add_paragraph("3. File layout", style="Heading 1")
    doc.add_paragraph(
        "A recording file is laid out as a 64-byte plaintext file header followed by "
        "a repeated sequence of plaintext batch header, encrypted payload, and "
        "plaintext trailer."
    )
    add_label_paragraph(
        doc,
        "Rev2 change",
        "The optional final summary record from the original design note is removed. "
        "Completed batches are the only durable recovery boundary."
    )

    doc.add_paragraph("4. File header", style="Heading 1")
    add_table(
        doc,
        ["Offset", "Size", "Field", "Rev2 value or meaning"],
        FILE_HEADER_ROWS,
        [720, 720, 2160, 5760],
    )

    doc.add_paragraph("5. Batch header", style="Heading 1")
    doc.add_paragraph(
        "Each batch begins with a 48-byte plaintext header. The header is written up "
        "front with nominal full-batch counts and may be rewritten with the actual "
        "counts if recording stops early."
    )
    add_table(
        doc,
        ["Offset", "Size", "Field", "Meaning"],
        BATCH_HEADER_ROWS,
        [720, 720, 2160, 5760],
    )

    doc.add_paragraph("6. Payload definition", style="Heading 1")
    add_label_paragraph(
        doc,
        "Payload dispatch",
        "payload_type is the only chunk-layout discriminator in rev2. The current "
        "implemented payload_type is 2, meaning one-second chunks of EMG ch1 plus IMU.",
    )
    add_label_paragraph(
        doc,
        "Chunk size",
        "One chunk equals 6400 bytes: int32 ch1[1000] plus ImuSample[100]. "
        "A full 10-second batch contains 10 chunks and therefore 64000 plaintext bytes.",
    )
    add_table(
        doc,
        ["Section", "Type", "Count", "Bytes", "Meaning"],
        PAYLOAD_ROWS,
        [1200, 1320, 960, 960, 4920],
    )
    add_label_paragraph(
        doc,
        "Lead-off policy",
        "Per-sample lead-off status is not stored in the binary payload. Instead, "
        "the firmware detects lead-off per 1-second chunk from the signal itself and "
        "reports the per-batch summary in the control CSV as ch1_leadoff_chunks.",
    )

    doc.add_paragraph("6.1 ImuSample layout", style="Heading 2")
    add_table(
        doc,
        ["Offset", "Type", "Fields"],
        IMU_ROWS,
        [900, 1620, 6840],
    )

    doc.add_paragraph("7. Encryption and nonce", style="Heading 1")
    add_label_paragraph(
        doc,
        "Cipher",
        "AES-256-GCM is the normative cipher for this format. Each batch is encrypted "
        "independently as one GCM stream over its plaintext payload.",
    )
    add_label_paragraph(
        doc,
        "Nonce",
        "The 96-bit batch nonce is nonce_prefix (8 bytes) plus batch_index encoded as "
        "a 32-bit little-endian value.",
    )
    add_label_paragraph(
        doc,
        "Rev2 nonce rationale",
        "nonce_prefix is a per-session entropy salt derived from ADS1292R analog noise, "
        "the free-running CPU cycle counter, and the device UID. This avoids dependence "
        "on RTC state or monotonic session numbers and therefore avoids nonce reuse risk "
        "under a fixed key.",
    )

    doc.add_paragraph("8. Batch trailer", style="Heading 1")
    add_table(
        doc,
        ["Offset", "Size", "Field", "Meaning"],
        [
            ("0", "4", "crc32_plaintext", "CRC32 of the plaintext payload"),
            ("4", "16", "tag", "AES-GCM authentication tag"),
            ("20", "4", "end_marker", '"ENDB"'),
        ],
        [720, 720, 2160, 5760],
    )
    add_label_paragraph(
        doc,
        "Recovery rule",
        "A batch missing the trailer or ENDB marker is treated as incomplete and "
        "discarded by readers. Scanning resumes at the next BATC marker.",
    )

    doc.add_paragraph("9. Secondary control CSV", style="Heading 1")
    doc.add_paragraph(
        "The control CSV contains one row per completed batch and no raw EMG or IMU "
        "measurement values."
    )
    add_table(
        doc,
        ["Column", "Meaning"],
        CSV_ROWS,
        [2520, 6840],
    )

    doc.add_paragraph("10. Physical scaling", style="Heading 1")
    add_label_paragraph(
        doc,
        "EMG scaling",
        "emg_uV = ch1 * emg_ref_mv * 1000.0 / ((1 << 23) * emg_pga_gain). "
        "With emg_ref_mv = 2420 and emg_pga_gain = 4, one count is approximately "
        "0.0721 microvolts.",
    )
    add_label_paragraph(doc, "Accelerometer", "g = raw / 16384 for the +/-2 g range.")
    add_label_paragraph(doc, "Gyroscope", "dps = raw / 131 for the +/-250 dps range.")
    add_label_paragraph(
        doc,
        "Magnetometer",
        "mx, my, and mz are stored as 18-bit unsigned values inside uint32 fields. "
        "A null field is 131072.",
    )

    doc.add_paragraph("11. Compatibility and versioning", style="Heading 1")
    add_label_paragraph(
        doc,
        "Current reader rule",
        "Readers must require version == 1 and payload_type == 2 for this rev2 "
        "document. Anything else is outside the current implemented format.",
    )
    add_label_paragraph(
        doc,
        "Legacy note",
        "Earlier development recordings are obsolete and incompatible. This includes "
        "pre-BATC files and early raw-status recordings that reused payload_type 2 "
        "with a different chunk size and a different nonce derivation.",
    )
    add_label_paragraph(
        doc,
        "Future changes",
        "If a future payload layout changes, payload_type should be bumped. If the "
        "header semantics change, version should be bumped.",
    )

    doc.add_paragraph("12. References", style="Heading 1")
    doc.add_paragraph(
        "This rev2 document aligns with the current implementation sources: the "
        "firmware writer in recording.c, acquisition support in acquisition.c, the "
        "desktop decoder decode_emgx.py, and the aligned markdown specification FORMAT.md."
    )

    return doc


def main():
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
